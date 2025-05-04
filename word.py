from docx import Document

# Création du document Word
doc = Document()
doc.add_heading("LYCÉE ACTION SOCIALE", level=1)
doc.add_paragraph("ANNÉE SCOLAIRE : 2024-2025")
doc.add_paragraph("CLASSE : 4ème A")
doc.add_paragraph("DATE : 12-04-2025")
doc.add_paragraph("PROPOSÉ PAR : Mr KABORE")
doc.add_paragraph("DURÉE : 2 Heures")
doc.add_heading("DEVOIR N°1 DE MATHÉMATIQUES (Troisième trimestre)", level=2)

# Exercice 1
doc.add_heading("EXERCICE 1 (7 pts)", level=3)
doc.add_paragraph("1) Simplifie les expressions suivantes : (1 pt × 2)")
doc.add_paragraph("U = AB⃗ + CD⃗ + FG⃗ + GC⃗ + DF⃗ + BA⃗")
doc.add_paragraph("V = AM⃗ + LK⃗ + MK⃗ + KM⃗ + KL⃗ + MA⃗ + CD⃗")

doc.add_paragraph("2) MNOP est un parallélogramme")
doc.add_paragraph("a) Écris plus simplement chacun des vecteurs suivants : (0,5 pt × 4)")
doc.add_paragraph("MP⃗ + MN⃗ = ")
doc.add_paragraph("PM⃗ + PO⃗ = ")
doc.add_paragraph("MN⃗ − MP⃗ = ")
doc.add_paragraph("OP⃗ + ON⃗ = ")

doc.add_paragraph("b) En utilisant le parallélogramme MNOP, démontre (utilise la relation de Chasles) que : (1,5 pt × 2)")
doc.add_paragraph("MO⃗ + NP⃗ = 2 MP⃗")
doc.add_paragraph("MO⃗ + PN⃗ = 2 MN⃗")

# Exercice 2
doc.add_heading("EXERCICE 2 (7 pts)", level=3)
doc.add_paragraph("Trace deux droites (D) et (D’) parallèles. (0,5 pt)")
doc.add_paragraph("1) Marque deux points I et J sur (D) puis K et L sur la droite (D’) tels que les bipoints (I, J) et (K, L) soient équipollents. (1 pt)")
doc.add_paragraph("2) Quelle est la nature du quadrilatère IJKL ? Justifie ta réponse. (1 pt)")
doc.add_paragraph("3) E est le symétrique de K par rapport à L. (0,5 pt)")
doc.add_paragraph("   F est le symétrique de J par rapport à I. (0,5 pt)")
doc.add_paragraph("a) Donne quatre bipoints équipollents sur la figure. (1 pt)")
doc.add_paragraph("b) Donne quatre vecteurs égaux sur la figure. (1 pt)")
doc.add_paragraph("c) Les bipoints (I, J), (L, F) et (K, J) sont-ils équipollents ? Justifie ta réponse. (1,5 pt)")

# Exercice 3
doc.add_heading("EXERCICE 3 (6 pts)", level=3)
doc.add_paragraph("Pendant une séance de cours, le professeur de Mathématiques d’une classe de 4ème au Lycée Action Sociale de Fada a mis ces informations et la figure ci-contre au tableau.")
doc.add_paragraph("ABCD est un parallélogramme et ED⃗ = DA⃗, BF⃗ = AB⃗.")
doc.add_paragraph("Appelé d’urgence à l’administration, il s’absente. C’est alors qu’un élève de la classe affirme que le point C est le milieu du segment [EF]. Les autres étonnés cherchent à vérifier en répondant aux consignes suivantes :")
doc.add_paragraph("1) Construis les points E et F. (1 pt)")
doc.add_paragraph("2) Justifie que EC⃗ = DB⃗. (2 pts)")
doc.add_paragraph("3) Justifie que CF⃗ = DB⃗. (2 pts)")
doc.add_paragraph("4) Dis si l’élève a raison ou pas. (1 pt)")

# Enregistrement du fichier
output_path = "DEV_TROIS_TRIMES_ACTION_SOCIALE_4e.docx"
doc.save(output_path)

output_path
